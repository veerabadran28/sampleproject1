import streamlit as st
import pandas as pd
import json
import plotly.graph_objects as go
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
from typing import Dict, Tuple, Optional, List
from io import StringIO
import numpy as np
import math
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from io import BytesIO
import pandas as pd
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

class DocumentGenerator:
    def __init__(self, common_config: Dict, s3_service, bedrock_service, doc_type: str = None):
        """Initialize document generator.
        
        Args:
            common_config: Common application configuration
            s3_service: S3 service instance
            bedrock_service: Bedrock service instance
            doc_type: Type of document to generate
        """
        self.common_config = common_config
        self.s3_service = s3_service
        self.bedrock_service = bedrock_service
        self.doc_type = doc_type
        self.config = self._load_config(doc_type)
        self.cached_data = {}

    #@st.cache_data(ttl=3600)
    def _load_config(_self, doc_type: str = None) -> Dict:
        """Load document generation configuration.
        
        Args:
            doc_type: Type of document to generate (e.g., 'ppnr')
            
        Returns:
            Dict containing merged configuration
        """
        try:
            # Load common config
            base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            #print(f"Base Path:{base_path}")
            config_dir = os.path.join(base_path, 'config')
            
            common_config_path = os.path.join(config_dir, 'doc_gen_config.json')
            if not os.path.exists(common_config_path):
                raise FileNotFoundError(f"Common config file not found at: {common_config_path}")
                
            with open(common_config_path, 'r', encoding='utf-8') as f:
                common_config = json.load(f)
                
            # If no doc_type specified, return only common config
            if not doc_type:
                return common_config
                
            # Validate document type
            if doc_type not in common_config.get('document_types', {}):
                raise ValueError(f"Invalid document type: {doc_type}")
                
            # Get document-specific config file name
            doc_config_file = common_config['document_types'][doc_type]['config_file']
            doc_config_path = os.path.join(config_dir, doc_config_file)
            
            if not os.path.exists(doc_config_path):
                raise FileNotFoundError(f"Document config file not found at: {doc_config_path}")
                
            # Load document-specific config
            with open(doc_config_path, 'r', encoding='utf-8') as f:
                doc_config = json.load(f)
                
            # Merge configurations
            # Common settings from doc_gen_config.json
            merged_config = {
                'document_type': doc_type,
                'input_settings': common_config['input_settings'],
                'output_settings': common_config['output_settings'],
                'chart_settings': common_config['chart_settings'],
                'document_settings': common_config['document_settings']
            }
            
            # Document specific settings override common settings
            merged_config.update({
                'document_info': doc_config['document_info'],
                'input_settings': {
                    **merged_config['input_settings'],
                    **doc_config['input_settings']
                },
                'output_settings': {
                    **merged_config['output_settings'],
                    **doc_config['output_settings']
                },
                'executive_summary': doc_config['executive_summary'],
                'business_sections': doc_config['business_sections']
            })
            
            # Verify narrative templates exist in each business section
            for section_key, section_data in merged_config['business_sections'].items():
                if 'narrative_templates' not in section_data:
                    raise ValueError(f"Narrative templates missing for business section: {section_key}")
                    
                required_templates = ['overall', 'baseline', 'stress']
                missing_templates = [t for t in required_templates if t not in section_data['narrative_templates']]
                if missing_templates:
                    raise ValueError(f"Missing required templates {missing_templates} for section {section_key}")
            
            return merged_config
            
        except Exception as e:
            st.error(f"Error loading configuration: {str(e)}")
            st.error(f"Current working directory: {os.getcwd()}")
            return {}

    def _process_csv_data(self, business_line: str, data_files: Dict[str, str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
        """Process CSV data for a specific business line.
        
        Args:
            business_line: Name of the business line to process
            data_files: Dictionary of scenario names and their file paths
            
        Returns:
            Tuple of DataFrames (assets_data, liabilities_data)
        """
        try:
            assets_data = []
            liabilities_data = []
            
            # Process base period (Q42023) data
            #print(data_files['base_period'])
            base_content = self.s3_service.get_document_content(data_files['base_period'])
            # Decode the content from bytes to a string
            try:
                base_content = base_content.decode('utf-8')
                #print(f"Here before base_content type: {type(base_content)}")
            except Exception as decode_error:
                print(f"Error decoding base_content: {decode_error}")
                base_content = None

            # Check if base_content is not empty
            if base_content:
                #print("Inside base_content")
                try:
                    # Use StringIO to simulate a file for pandas
                    csv_buffer = StringIO(base_content)
                    base_df = pd.read_csv(csv_buffer, delimiter='|')  # Adjust the delimiter if needed
                    #print("Before base_df")
                    #print(base_df.head())  # Display the DataFrame for debugging
                    # Filter for the specific business line
                    print(f"business_line:[{business_line}]")
                    base_df = base_df[base_df['business'] == business_line]
                    #print("After base_df")
                    # Get 2023Q4 values
                    q4_2023_assets = round(base_df[base_df['level_2'] == 'U10000000']['balance'].sum() / 100000000, 1)  # Convert to billions
                    q4_2023_liabilities = round(base_df[base_df['level_2'] == 'U20000000']['balance'].sum() / 100000000, 1)
                    #print(f"q4_2023_assets:[{q4_2023_assets}], q4_2023_liabilities:[{q4_2023_liabilities}]")
                    
                    # Create initial rows for each scenario
                    scenarios = ['BHC Base', 'BHC Stress', 'Sup Base', 'Sup Sev Adv']
                    for scenario in scenarios:
                        assets_data.append({'Scenario ($bn)': scenario, '2023Q4': q4_2023_assets})
                        liabilities_data.append({'Scenario ($bn)': scenario, '2023Q4': -q4_2023_liabilities})  # Negative for liabilities
                    
                    #print("assets_data")
                    #print(assets_data)
                    #print("liabilities_data")
                    #print(liabilities_data)
                except Exception as read_error:
                    print(f"Error read_csv: {str(read_error)}")
            else:
                print("base_content is empty or invalid")
                raise Exception
            
            # Process scenario data
            for scenario, file_path in data_files['scenarios'].items():
                content = self.s3_service.get_document_content(file_path)
                
                content = content.decode('utf-8')
                csv_buffer = StringIO(content)
                
                if content:
                    df = pd.read_csv(csv_buffer, delimiter='|')
                    df = df[df['business'] == business_line]
                    
                    # Process each quarter
                    quarters = ['2024Q1', '2024Q2', '2024Q3', '2024Q4', 
                              '2025Q1', '2025Q2', '2025Q3', '2025Q4', '2026Q1']
                    
                    for quarter in quarters:
                        #print(f"scenario data for:{scenario}")
                        #print(df.head())
                        quarter_df = df[df['period_id'] == quarter]
                        assets = round(quarter_df[quarter_df['level_2'] == 'U10000000']['balance'].sum() / 100000000, 1)
                        liabilities = round(quarter_df[quarter_df['level_2'] == 'U20000000']['balance'].sum() / 100000000, 1)
                        #print(f"assets:[{assets}], liabilities:[{liabilities}]")
                        # Update scenario rows
                        try:
                            assets_idx = next(i for i, d in enumerate(assets_data) if d['Scenario ($bn)'] == scenario)
                        except Exception as e:
                            print(f"Error assets_idx: {str(e)}")
                        #print(f"assets_idx:[{assets_idx}]")
                        liabilities_idx = next(i for i, d in enumerate(liabilities_data) if d['Scenario ($bn)'] == scenario)
                        #print(f"liabilities_idx:[{liabilities_idx}]")
                        
                        assets_data[assets_idx][quarter] = assets
                        liabilities_data[liabilities_idx][quarter] = -liabilities  # Negative for liabilities
                        #print("assets_data")
                        #print(assets_data)
                        #print("liabilities_data")
                        #print(liabilities_data)
            
            # Convert to DataFrames
            assets_df = pd.DataFrame(assets_data)
            liabilities_df = pd.DataFrame(liabilities_data)
            
            # Calculate 9Q Average
            quarter_cols = ['2024Q1', '2024Q2', '2024Q3', '2024Q4', 
                          '2025Q1', '2025Q2', '2025Q3', '2025Q4', '2026Q1']
            
            assets_df['9Q Avg'] = assets_df[quarter_cols].mean(axis=1).round(1)
            liabilities_df['9Q Avg'] = liabilities_df[quarter_cols].mean(axis=1).round(1)
            
            #print("assets_df1")
            #print(assets_df)
            #print("liabilities_df1")
            #print(liabilities_df)
            
            # Calculate Dev. from BHC Base
            bhc_base_assets = assets_df[assets_df['Scenario ($bn)'] == 'BHC Base']['9Q Avg'].values[0]
            bhc_base_liabilities = liabilities_df[liabilities_df['Scenario ($bn)'] == 'BHC Base']['9Q Avg'].values[0]
            
            def calc_deviation(row, base_value):
                if row['Scenario ($bn)'] == 'BHC Base':
                    return None
                return ((row['9Q Avg'] / base_value) - 1) * 100 if base_value != 0 else None
            
            assets_df['Dev. from BHC Base'] = assets_df.apply(lambda row: calc_deviation(row, bhc_base_assets), axis=1)
            liabilities_df['Dev. from BHC Base'] = liabilities_df.apply(lambda row: calc_deviation(row, bhc_base_liabilities), axis=1)
            
            #print("assets_df2")
            #print(assets_df)
            #print("liabilities_df2")
            #print(liabilities_df)
            
            # Format the deviation values
            def format_deviation(val):
                if pd.isna(val):
                    return ''
                return f"({abs(val):.1f})%" if val < 0 else f"{val:.1f}%"
            
            assets_df['Dev. from BHC Base'] = assets_df['Dev. from BHC Base'].apply(
                lambda x: format_deviation(x) if pd.notnull(x) else '')
            liabilities_df['Dev. from BHC Base'] = liabilities_df['Dev. from BHC Base'].apply(
                lambda x: format_deviation(x) if pd.notnull(x) else '')
            
            #print("assets_df3")
            #print(assets_df)
            #print("liabilities_df3")
            #print(liabilities_df)
            
            return assets_df, liabilities_df

        except Exception as e:
            st.error(f"Error processing CSV data: {str(e)}")
            return pd.DataFrame(), pd.DataFrame()

    def save_cache(self) -> bool:
        """Save the current cache to S3."""
        try:
            if not self.cached_data:
                return False
                
            # Define json filename
            json_filename = "2024_ppnr_methodology_and_process_overview.json"
            
            # Save locally first
            with open(json_filename, 'w') as f:
                json.dump(self.cached_data, f, indent=2)
                
            # Upload to S3
            s3_path = os.path.join(
                self.config['output_settings']['s3_output_path'],
                json_filename
            ).replace('\\', '/')
            
            success = self.s3_service.upload_document(json_filename, s3_path)
            
            # Clean up local file
            try:
                os.remove(json_filename)
            except:
                pass
                
            return success
            
        except Exception as e:
            st.error(f"Error saving cache: {str(e)}")
            return False
    
    def generate_driver_table(self, driver_config: Dict) -> go.Figure:
        """Generate driver information table."""
        try:
            # Create DataFrame from driver configuration
            df = pd.DataFrame([{
                'Dependent Variable': driver_config['dependent_var'],
                'Independent Variable': driver_config['independent_var'],
                'Lag': driver_config['lag'],
                'Direction': driver_config['direction']
            }])

            # Create Plotly table with minimal height
            fig = go.Figure(data=[go.Table(
                header=dict(
                    values=list(df.columns),
                    fill_color='#0051A2',  # Blue background
                    font=dict(color='white', size=12),
                    align='center',
                    height=30  # Reduced header height
                ),
                cells=dict(
                    values=[df[col] for col in df.columns],
                    fill_color='white',
                    align='center',
                    height=45  # Reduced cell height
                )
            )])

            # Update layout to remove extra space
            fig.update_layout(
                margin=dict(l=0, r=0, t=0, b=0),
                height=85  # Total height = header height + cell height
            )
            
            return fig

        except Exception as e:
            st.error(f"Error generating driver table: {str(e)}")
            return None

    def generate_line_chart(self, data: pd.DataFrame, title: str, y_label: str) -> Optional[go.Figure]:
        """Generate line chart for data visualization."""
        try:
            print("Starting chart generation...")
            print(f"Input data shape: {data.shape}")
            print(f"Columns: {data.columns.tolist()}")
            
            fig = go.Figure()
            
            # Define consistent colors and order for scenarios
            scenario_colors = {
                'BHC Base': '#000000',    # Black
                'BHC Stress': '#0000FF',  # Blue
                'Sup Base': '#006400',    # Dark Green
                'Sup Sev Adv': '#800080'  # Purple
            }
            
            # Get quarters and sort them
            quarters = [col for col in data.columns if col.startswith(('20')) 
                    and col not in ['Scenario ($bn)', '9Q Avg', 'Dev. from BHC Base']]
            quarters.sort()
            print(f"Quarters found: {quarters}")
            
            # Add traces in specific order
            for scenario in scenario_colors.keys():
                print(f"Processing scenario: {scenario}")
                if scenario in data['Scenario ($bn)'].values:
                    scenario_data = data[data['Scenario ($bn)'] == scenario]
                    print(f"Found data for scenario: {scenario}")
                    
                    y_values = []
                    for q in quarters:
                        try:
                            val = scenario_data[q].values[0]
                            val = round(float(val), 1)
                            y_values.append(val)
                        except Exception as val_error:
                            print(f"Error processing value for {q}: {str(val_error)}")
                            y_values.append(None)
                    
                    print(f"Y values for {scenario}: {y_values}")
                    
                    fig.add_trace(go.Scatter(
                        x=quarters,
                        y=y_values,
                        name=scenario,
                        line=dict(
                            color=scenario_colors[scenario],
                            width=1.5
                        ),
                        mode='lines+markers',
                        marker=dict(
                            size=5,
                            symbol='circle'
                        )
                    ))
            
            # Update layout
            fig.update_layout(
                title=dict(
                    text=title,
                    x=0.5,
                    y=0.95,
                    xanchor='center',
                    yanchor='top',
                    font=dict(size=14)
                ),
                xaxis=dict(
                    title="",
                    showgrid=True,
                    gridcolor='lightgray',
                    gridwidth=1,
                    tickmode='array',
                    ticktext=[q.replace('20', '') for q in quarters],
                    tickvals=quarters,
                    tickangle=0,
                    showline=True,
                    linewidth=1,
                    linecolor='black'
                ),
                yaxis=dict(
                    title=dict(
                        text=y_label,
                        font=dict(size=12)
                    ),
                    showgrid=True,
                    gridcolor='lightgray',
                    gridwidth=1,
                    showline=True,
                    linewidth=1,
                    linecolor='black'
                ),
                showlegend=True,
                legend=dict(
                    yanchor="top",
                    y=0.99,
                    xanchor="right",
                    x=0.99,
                    bgcolor='rgba(255, 255, 255, 0.9)',
                    bordercolor='black',
                    borderwidth=1
                ),
                plot_bgcolor='white',
                width=800,
                height=400,
                margin=dict(l=50, r=50, t=50, b=50, pad=4)
            )
            
            print("Chart generated successfully")
            return fig
            
        except Exception as e:
            print(f"Error generating line chart: {str(e)}")
            print(f"Data columns: {data.columns}")
            st.error(f"Error generating line chart: {str(e)}")
            return None

    def _format_data_for_narrative(self, assets_df: pd.DataFrame, liabilities_df: pd.DataFrame) -> str:
        """Format data for narrative generation."""
        try:
            formatted_data = "Assets Projections:\n"
            formatted_data += assets_df.to_string()
            formatted_data += "\n\nLiabilities Projections:\n"
            formatted_data += liabilities_df.to_string()
            return formatted_data
        except Exception as e:
            st.error(f"Error formatting data for narrative: {str(e)}")
            return ""
        
    def generate_narratives(self, section_key: str, data: Dict[str, pd.DataFrame], driver_data: Dict) -> Dict[str, str]:
        """Generate narratives using the Nova model."""
        try:
            narratives = {}
            
            # Get business-specific narrative templates from config
            if section_key not in self.config['business_sections']:
                raise ValueError(f"Invalid section key: {section_key}")
                
            section_config = self.config['business_sections'][section_key]
            if 'narrative_templates' not in section_config:
                raise ValueError(f"No narrative templates found for section: {section_key}")
                
            narrative_templates = section_config['narrative_templates']
            
            # Format data for prompt
            formatted_data = self._format_data_for_narrative(
                data['assets'],
                data['liabilities']
            )
            print("formatted_data")
            print(formatted_data)
            # Generate overall narrative
            overall_prompt = self._create_narrative_prompt(
                narrative_templates['overall'],
                formatted_data,
                section_key,
                'overall',
                driver_data
            )
            print("overall_prompt")
            print(overall_prompt)
            narratives['overall'] = self.bedrock_service.invoke_nova_model(overall_prompt)
            #narratives['overall'] = self._format_model_response(response)
            
            # Generate baseline narrative
            baseline_prompt = self._create_narrative_prompt(
                narrative_templates['baseline'],
                formatted_data,
                section_key,
                'baseline',
                driver_data
            )
            #response = self.bedrock_service.invoke_nova_model(baseline_prompt)
            #narratives['baseline'] = self._format_model_response(response)
            narratives['baseline'] = self.bedrock_service.invoke_nova_model(baseline_prompt)
            
            # Generate stress narrative
            stress_prompt = self._create_narrative_prompt(
                narrative_templates['stress'],
                formatted_data,
                section_key,
                'stress',
                driver_data
            )
            #response = self.bedrock_service.invoke_nova_model(stress_prompt)
            #narratives['stress'] = self._format_model_response(response)
            narratives['stress'] = self.bedrock_service.invoke_nova_model(stress_prompt)
            
            return narratives

        except Exception as e:
            st.error(f"Error generating narratives: {str(e)}")
            return {}

    def _create_narrative_prompt(self, template: Dict, data: str, section_key: str, narrative_type: str, driver_data: pd.DataFrame) -> str:
        """Create specific narrative prompt based on type."""
        try:
            # Get business-specific context if available
            section_config = self.config['business_sections'][section_key]
            business_context = section_config.get('business_context', '')
            
            prompt = f"""Please analyze the provided financial projections and driver information to generate detailed CCAR model narratives as per the given example output.

Data to Analyze:
{data}

Drivers: The following driver information is relevant for scenario analysis and narrative generation.
{driver_data}

Example Output provided below: This is how the CCAR users write narratives for the overview business section and baseline, and stress scenario sections. Carefully review and adhere to the same format provided in the example with no deviation.
{template['example']}

Additional Guidelines:
1. Strictly follow the exact formatting shown in the example output with no deviation
2. Focus on business-specific risk factors, drivers and considerations

Strictly provide only the {narrative_type} narrative following exactly the format shown in the Example Output above."""

            return prompt

        except Exception as e:
            st.error(f"Error creating narrative prompt: {str(e)}")
            return ""    
    
    def _format_model_response(self, response_text: str) -> str:
        """
        Format the model response text to properly display in Streamlit UI.
        Handles special characters and currency symbols.
        
        Args:
            response_text: The original response text from the model
            
        Returns:
            Formatted text safe for Streamlit display
        """
        # Handle currency amounts with $ symbol
        import re
        
        def replace_currency(match):
            # Get the full match
            full_match = match.group(0)
            
            # If it's $(number), format it with a space after $
            if full_match.startswith('$('):
                return r'\\$(' + match.group(1) + ')'
            
            # For regular $number
            return r'\\$' + match.group(1)
        
        # Replace $number or $(number) patterns
        # This handles both positive and negative currency values
        formatted_text = re.sub(r'\$(\([0-9.,]+\)|\d*\.?\d+)', replace_currency, response_text)
        
        # Additional formatting if needed
        formatted_text = formatted_text.replace('$bn', r'\\$bn')  # Handle "$bn" specifically
        
        return formatted_text
    
    def generate_document(self) -> Tuple[Optional[str], Optional[str]]:
        """Generate document with all content."""
        try:
            # Define filenames
            base_filename = "2024_ppnr_methodology_and_process_overview"
            docx_filename = f"{base_filename}.docx"
            json_filename = f"{base_filename}.json"
            
            # Get base path from config
            s3_base_path = self.config['output_settings']['s3_output_path']

            # Generate content if not already cached
            if not self.cached_data:
                self.cached_data = self._generate_complete_cache()

            # Create Word document
            doc = Document()
            self._add_document_content(doc)

            # Save files locally first
            with open(json_filename, 'w') as f:
                json.dump(self.cached_data, f, indent=2)
                
            doc.save(docx_filename)            

            # Upload to S3
            try:
                for filename in [docx_filename, json_filename]:
                    full_path = os.path.join(s3_base_path, filename).replace('\\', '/')
                    self.s3_service.upload_document(filename, full_path)
                        
            except Exception as e:
                st.error(f"Error uploading files to S3: {str(e)}")
                return None, None
            finally:
                # Clean up local files
                for filename in [docx_filename, json_filename]:
                    try:
                        os.remove(filename)
                    except:
                        pass

            # Return the full S3 paths
            docx_path = os.path.join(s3_base_path, docx_filename).replace('\\', '/')
            return docx_path

        except Exception as e:
            st.error(f"Error generating document: {str(e)}")
            return None, None

    def load_latest_cache(self) -> bool:
        """Load the cached data from S3."""
        try:
            # Construct the full S3 path
            json_file = "2024_ppnr_methodology_and_process_overview.json"
            s3_path = os.path.join(
                self.config['output_settings']['s3_output_path'],
                json_file
            ).replace('\\', '/')
            
            # Check if file exists before trying to load it
            if not self.s3_service.file_exists(s3_path):
                return False
            
            # Load the content
            content = self.s3_service.get_document_content(s3_path)
            if content:
                self.cached_data = json.loads(content.decode('utf-8'))
                return True
                
            return False
            
        except Exception as e:
            st.error(f"Error loading cached data: {str(e)}")
            return False    
    
    def _create_section_cache(self, section_key: str, section_config: Dict) -> Dict:
        """Create cache data for a section."""
        try:
            cache_data = {
                'title': section_config['title'],
                'topics': {}
            }
            
            # Process each topic
            for topic_key, topic_config in section_config['topics'].items():
                topic_data = {
                    'title': topic_config['title'],
                    'driver_info': topic_config.get('driver_info', {}),
                    'projections': {},
                    'narratives': {}
                }
                
                # Process data
                data = self._process_csv_data(
                    business_line=section_config['title'],
                    data_files={
                        'base_period': self.config['input_settings']['base_period_file'],
                        'scenarios': self.config['input_settings']['scenario_files']
                    }
                )
                
                if data and len(data) == 2:
                    assets_df, liabilities_df = data
                    
                    # Store projections data
                    topic_data['projections'] = {
                        'assets': assets_df.to_dict('records') if not assets_df.empty else None,
                        'liabilities': liabilities_df.to_dict('records') if not liabilities_df.empty else None
                    }
                    
                    # Generate and store narratives
                    narratives = self.generate_narratives(
                        section_key,
                        {'assets': assets_df, 'liabilities': liabilities_df}
                    )
                    topic_data['narratives'] = narratives
                
                cache_data['topics'][topic_key] = topic_data
                
            return cache_data
            
        except Exception as e:
            st.error(f"Error creating section cache: {str(e)}")
            return {}
    
    def _generate_complete_cache(self) -> Dict:
        """Generate complete cache data for all sections."""
        cache_data = {
            'document_info': self.config['document_info'],
            'executive_summary': self.config['executive_summary'],
            'business_sections': {}
        }
        
        # Process each business section
        for section_key, section_config in self.config['business_sections'].items():
            cache_data['business_sections'][section_key] = self._create_section_cache(
                section_key, section_config)
                
        return cache_data

    def get_latest_documents(self) -> Optional[Dict]:
        """Get the existing document versions from S3."""
        try:
            # Get S3 base path from config
            s3_base_path = self.config['output_settings']['s3_output_path']
            docs = {}
            
            base_filename = "2024_ppnr_methodology_and_process_overview"
            for ext in ['.docx', '.json']:
                filename = f"{base_filename}{ext}"
                full_path = os.path.join(s3_base_path, filename).replace('\\', '/')
                
                # Check if file exists
                if self.s3_service.file_exists(full_path):
                    docs[ext[1:]] = {
                        'filename': filename,
                        'url': full_path
                    }
            
            return docs if docs else None

        except Exception as e:
            st.error(f"Error retrieving documents: {str(e)}")
            return None
    
    def _validate_sections(self, cached_data: Dict) -> Tuple[bool, List[str]]:
        """Validate that all required sections are present and have content.
        
        Returns:
            Tuple[bool, List[str]]: (is_valid, list of missing/incomplete sections)
        """
        missing_sections = []
        
        try:            
            # Check executive summary
            if 'executive_summary' not in cached_data:
                missing_sections.append("Executive Summary section is missing")
            elif not cached_data['executive_summary'].get('sections'):
                missing_sections.append("Executive Summary content is incomplete")
                
            # Check business sections
            if 'business_sections' not in cached_data:
                missing_sections.append("Business sections are missing")
                return False, missing_sections
                
            # For each business section, verify required components
            for section_key, section_data in cached_data['business_sections'].items():
                # Check section has topics
                if 'topics' not in section_data:
                    missing_sections.append(f"Section '{section_key}' is missing topics")
                    continue
                
                if 'narrative_templates' not in self.config['business_sections'][section_key]:
                    missing_sections.append(f"Section '{section_key}' is missing narrative templates")
                    continue
                    
                for topic_key, topic_data in section_data['topics'].items():
                    # Check required topic components
                    missing_components = []
                    for required in ['title', 'projections', 'narratives']:
                        if required not in topic_data:
                            missing_components.append(required)
                    
                    if missing_components:
                        missing_sections.append(
                            f"Topic '{topic_key}' in section '{section_key}' is missing: {', '.join(missing_components)}"
                        )
                        
                    # Verify projections data exists
                    projections = topic_data.get('projections', {})
                    if not projections:
                        missing_sections.append(f"No projections data for '{section_key}/{topic_key}'")
                    else:
                        for proj_type in ['assets', 'liabilities']:
                            if proj_type not in projections:
                                missing_sections.append(
                                    f"Missing {proj_type} projections for '{section_key}/{topic_key}'"
                                )
                        
                    # Verify narratives exist
                    narratives = topic_data.get('narratives', {})
                    if not narratives:
                        missing_sections.append(f"No narratives for '{section_key}/{topic_key}'")
                    else:
                        for narr_type in ['overall', 'baseline', 'stress']:
                            if narr_type not in narratives:
                                missing_sections.append(
                                    f"Missing {narr_type} narrative for '{section_key}/{topic_key}'"
                                )
            
            is_valid = len(missing_sections) == 0
            return is_valid, missing_sections
            
        except Exception as e:
            missing_sections.append(f"Error during validation: {str(e)}")
            return False, missing_sections
    
    def _set_cell_background(self, cell, hex_color: str):
        """Set cell background color."""
        cell_properties = cell._tc.get_or_add_tcPr()
        shading_element = OxmlElement('w:shd')
        shading_element.set(qn('w:fill'), hex_color)
        cell_properties.append(shading_element)
    
    def _add_document_content(self, doc: Document):
        """Add content to the Word document from cached data."""
        try:
            print("Starting document generation...")
            
            # Set document properties
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            
            # Add title
            title = doc.add_heading(self.cached_data['document_info']['title'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Add spacing

            # Add Executive Summary
            if 'executive_summary' in self.cached_data:
                print("Adding executive summary...")
                doc.add_heading('Executive Summary', level=1)
                
                for section in self.cached_data['executive_summary']['sections']:
                    # Add section heading
                    doc.add_heading(section['title'], level=2)
                    para = doc.add_paragraph()
                    para.add_run(section['content'])
                    doc.add_paragraph()  # Add spacing
                    
                doc.add_page_break()

            # Add business sections from cached data
            if 'business_sections' in self.cached_data:
                print("Adding business sections...")
                for section_key, section_data in self.cached_data['business_sections'].items():
                    # Add section heading
                    doc.add_heading(section_data['title'], level=1)
                    
                    # Process each topic
                    for topic_key, topic_data in section_data['topics'].items():
                        print(f"Processing topic: {topic_key}")
                        
                        # Add topic heading
                        doc.add_heading(topic_data['title'], level=2)
                        
                        # Add driver information
                        if 'driver_info' in topic_data and 'table' in topic_data['driver_info']:
                            doc.add_heading('Driver Information', level=3)
                            driver_info = topic_data['driver_info']['table']
                            
                            # Create table
                            table = doc.add_table(rows=2, cols=4)
                            table.style = 'Table Grid'
                            
                            # Add headers
                            headers = ['Dependent Variable', 'Independent Variable', 'Lag', 'Direction']
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(headers):
                                cell = header_cells[i]
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = cell.paragraphs[0].add_run(header)
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                # Set blue background
                                self._set_cell_background(cell, "0051A2")
                            
                            # Add data
                            data = [
                                driver_info['dependent_var'],
                                driver_info['independent_var'],
                                str(driver_info['lag']),
                                driver_info['direction']
                            ]
                            data_cells = table.rows[1].cells
                            for i, value in enumerate(data):
                                cell = data_cells[i]
                                cell.text = str(value)
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            doc.add_paragraph()  # Add spacing
                        
                        # Add projections
                        if 'projections' in topic_data:
                            # Assets projections
                            if 'assets' in topic_data['projections']:
                                assets_data = topic_data['projections']['assets']
                                if assets_data:
                                    doc.add_heading("FI Financing total assets projections", level=3)
                                    assets_df = pd.DataFrame(assets_data)
                                    
                                    # Create table
                                    table = doc.add_table(rows=len(assets_df)+1, cols=len(assets_df.columns))
                                    table.style = 'Table Grid'
                                    
                                    # Add headers
                                    header_cells = table.rows[0].cells
                                    for j, col in enumerate(assets_df.columns):
                                        cell = header_cells[j]
                                        cell.text = str(col)
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = cell.paragraphs[0].runs[0]
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        # Set blue background
                                        self._set_cell_background(cell, "0051A2")
                                    
                                    # Add data
                                    for i, row in assets_df.iterrows():
                                        for j, value in enumerate(row):
                                            cell = table.rows[i+1].cells[j]
                                            cell.text = str(value)
                                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    doc.add_paragraph()  # Add spacing
                                    
                                    # Add chart if available
                                    try:
                                        fig = self.generate_line_chart(
                                            assets_df,
                                            "FI Financing total Consolidated assets",
                                            "Assets ($bn)"
                                        )
                                        if fig:
                                            print("Before calling _create_chart_image")
                                            img_bytes = self._create_chart_image(fig)
                                            print("After calling _create_chart_image")
                                            if img_bytes:
                                                doc.add_picture(BytesIO(img_bytes), width=Inches(6))
                                                doc.add_paragraph()
                                    except Exception as chart_error:
                                        print(f"Error adding assets chart: {str(chart_error)}")
                            
                            # Liabilities projections
                            if 'liabilities' in topic_data['projections']:
                                liabilities_data = topic_data['projections']['liabilities']
                                if liabilities_data:
                                    doc.add_heading("FI Financing total liabilities projections", level=3)
                                    liabilities_df = pd.DataFrame(liabilities_data)
                                    
                                    # Create table
                                    table = doc.add_table(rows=len(liabilities_df)+1, cols=len(liabilities_df.columns))
                                    table.style = 'Table Grid'
                                    
                                    # Add headers
                                    header_cells = table.rows[0].cells
                                    for j, col in enumerate(liabilities_df.columns):
                                        cell = header_cells[j]
                                        cell.text = str(col)
                                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = cell.paragraphs[0].runs[0]
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        # Set blue background
                                        self._set_cell_background(cell, "0051A2")
                                    
                                    # Add data
                                    for i, row in liabilities_df.iterrows():
                                        for j, value in enumerate(row):
                                            cell = table.rows[i+1].cells[j]
                                            cell.text = str(value)
                                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    doc.add_paragraph()  # Add spacing
                                    
                                    # Add chart if available
                                    try:
                                        fig = self.generate_line_chart(
                                            liabilities_df,
                                            "FI Financing total Consolidated liabilities",
                                            "Liabilities ($bn)"
                                        )
                                        if fig:
                                            img_bytes = self._create_chart_image(fig)
                                            if img_bytes:
                                                doc.add_picture(BytesIO(img_bytes), width=Inches(6))
                                                doc.add_paragraph()
                                    except Exception as chart_error:
                                        print(f"Error adding liabilities chart: {str(chart_error)}")
                        
                        # Add narratives
                        if 'narratives' in topic_data:
                            narratives = topic_data['narratives']
                            
                            if 'overall' in narratives:
                                doc.add_heading('Analysis', level=2)
                                doc.add_paragraph(narratives['overall'])
                                doc.add_paragraph()  # Add spacing
                            
                            if 'baseline' in narratives:
                                doc.add_heading('Baseline Scenario', level=3)
                                doc.add_paragraph(narratives['baseline'])
                                doc.add_paragraph()  # Add spacing
                            
                            if 'stress' in narratives:
                                doc.add_heading('Stress Scenarios', level=3)
                                doc.add_paragraph(narratives['stress'])
                                doc.add_paragraph()  # Add spacing
                    
                    doc.add_page_break()  # Add page break after each section

            print("Document content added successfully")

        except Exception as e:
            print(f"Error adding document content: {str(e)}")
            import traceback
            print(f"Full error: {traceback.format_exc()}")
            raise

    def _add_section_content(self, doc: Document, section_key: str, section_data: Dict):
        """Add section content to document."""
        try:
            print(f"Adding content for section: {section_key}")
            # Add section heading
            doc.add_heading(section_data['title'], level=1)

            # Process topics
            if 'topics' in section_data:
                for topic_key, topic_data in section_data['topics'].items():
                    print(f"Processing topic: {topic_key}")
                    # Add topic heading
                    doc.add_heading(topic_data['title'], level=2)

                    # Add driver information if available
                    self._add_driver_info(doc, topic_data)
                    
                    # Get cached data
                    cached_topic = self.cached_data.get('business_sections', {}).get(section_key, {}).get('topics', {}).get(topic_key, {})

                    # Add projections
                    self._add_projections_content(doc, topic_data, cached_topic)

                    # Add narratives
                    self._add_narratives_content(doc, cached_topic)

        except Exception as e:
            print(f"Error in _add_section_content: {str(e)}")
            st.error(f"Error adding section content: {str(e)}")
            raise

    def _add_driver_info(self, doc: Document, topic_data: Dict):
        """Add driver information table to document."""
        try:
            if 'driver_info' in topic_data:
                print("Adding driver information...")
                doc.add_heading('Driver Information', level=2)
                driver_table = doc.add_table(rows=2, cols=4)
                driver_info = topic_data['driver_info']['table']
                
                # Add headers
                headers = ['Dependent Variable', 'Independent Variable', 'Lag', 'Direction']
                header_row = driver_table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add data
                data_row = driver_table.rows[1]
                data = [
                    driver_info['dependent_var'],
                    driver_info['independent_var'],
                    str(driver_info['lag']),
                    driver_info['direction']
                ]
                
                for i, value in enumerate(data):
                    cell = data_row.cells[i]
                    cell.text = value
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Add spacing
                print("Driver information added successfully")
        except Exception as e:
            print(f"Error adding driver information: {str(e)}")

    def _add_projections_content(self, doc: Document, topic_data: Dict, cached_topic: Dict):
        """Add projections content including tables and charts."""
        try:
            if 'projections' in topic_data and 'projections' in cached_topic:
                print("Adding projections...")
                
                # Handle assets projections
                self._add_projection_type(
                    doc,
                    topic_data,
                    cached_topic,
                    'assets',
                    'Assets ($bn)'
                )
                
                # Handle liabilities projections
                self._add_projection_type(
                    doc,
                    topic_data,
                    cached_topic,
                    'liabilities',
                    'Liabilities ($bn)'
                )
                print("Projections added successfully")
        except Exception as e:
            print(f"Error adding projections content: {str(e)}")

    def _add_projection_type(self, doc: Document, topic_data: Dict, cached_topic: Dict, proj_type: str, y_label: str):
        """Add specific projection type (assets or liabilities) content."""
        try:
            if proj_type in topic_data['projections']:
                print(f"Processing {proj_type} projections...")
                doc.add_heading(topic_data['projections'][proj_type]['title'], level=3)
                
                projection_data = cached_topic['projections'][proj_type]
                if projection_data:
                    print(f"Creating {proj_type} table...")
                    df = pd.DataFrame(projection_data)
                    
                    # Add table
                    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
                    
                    # Add headers
                    for j, col in enumerate(df.columns):
                        cell = table.cell(0, j)
                        cell.text = str(col)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add data
                    for i, row in df.iterrows():
                        for j, value in enumerate(row):
                            cell = table.cell(i+1, j)
                            cell.text = str(value)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()  # Add spacing
                    print(f"{proj_type} table added")

                    # Add chart
                    print(f"Adding {proj_type} chart...")
                    chart_title = topic_data['projections'][proj_type]['chart_title']
                    fig = self.generate_line_chart(df, chart_title, y_label)
                    print(f"After adding {proj_type} chart...")
                    if fig:
                        print(f"Before creating {proj_type} chart image...")
                        img_bytes = self._create_chart_image(fig)
                        print(f"After creating {proj_type} chart image...")
                        if img_bytes:
                            try:
                                img_stream = BytesIO(img_bytes)
                                print(f"Before adding {proj_type} picture...")
                                doc.add_picture(img_stream, width=Inches(6))
                                print(f"After adding {proj_type} picture...")
                                doc.add_paragraph()
                                print(f"{proj_type} chart added successfully")
                            except Exception as chart_error:
                                print(f"Error adding {proj_type} chart: {str(chart_error)}")
                                doc.add_paragraph(f"[Chart: {chart_title}]")
                        else:
                            print(f"Failed to create {proj_type} chart image")
                            doc.add_paragraph(f"[Chart: {chart_title}]")
        except Exception as e:
            print(f"Error adding {proj_type} projection: {str(e)}")

    def _add_narratives_content(self, doc: Document, cached_topic: Dict):
        """Add narrative sections to document."""
        try:
            if 'narratives' in cached_topic:
                print("Adding narratives...")
                narratives = cached_topic['narratives']
                
                if 'overall' in narratives:
                    doc.add_heading('Overall Analysis', level=3)
                    doc.add_paragraph(narratives['overall'])
                
                if 'baseline' in narratives:
                    doc.add_heading('Baseline Scenario Analysis', level=3)
                    doc.add_paragraph(narratives['baseline'])
                
                if 'stress' in narratives:
                    doc.add_heading('Stress Scenarios Analysis', level=3)
                    doc.add_paragraph(narratives['stress'])
                    
                print("Narratives added successfully")
        except Exception as e:
            print(f"Error adding narratives: {str(e)}")
    
    def _save_chart_to_document(self, doc: Document, fig: go.Figure, title: str) -> bool:
        """Save chart to document using a simpler approach."""
        try:
            print("Starting chart save process...")
            
            # First convert to PNG bytes
            print("Converting to PNG...")
            try:
                img_bytes = fig.to_image(
                    format="png",
                    width=800,
                    height=400,
                    scale=1.0,
                    engine="kaleido"
                )
                print("Successfully converted to PNG")
            except Exception as png_error:
                print(f"Error converting to PNG: {str(png_error)}")
                # Try with lower quality
                try:
                    print("Trying with lower quality...")
                    img_bytes = fig.to_image(
                        format="png",
                        width=400,
                        height=300,
                        scale=0.5,
                        engine="kaleido"
                    )
                except Exception as retry_error:
                    print(f"Error on retry: {str(retry_error)}")
                    return False

            # Create BytesIO object
            print("Creating BytesIO object...")
            img_stream = BytesIO(img_bytes)
            print("BytesIO object created")

            # Add to document
            print("Adding to document...")
            try:
                doc.add_picture(img_stream, width=Inches(6))
                doc.add_paragraph()  # Add spacing
                print("Added to document successfully")
                return True
            except Exception as doc_error:
                print(f"Error adding to document: {str(doc_error)}")
                return False

        except Exception as e:
            print(f"General error in chart save: {str(e)}")
            return False

    def _create_chart_image(self, fig: go.Figure) -> Optional[bytes]:
        """Create chart image bytes with simplified settings."""
        try:
            print("Inside _create_chart_image")
            # Use minimal layout settings
            fig.update_layout(
                width=800,
                height=400,
                margin=dict(l=50, r=50, t=50, b=50),
                plot_bgcolor='white',
                showlegend=True
            )
            print("Updated layout")
            # Convert to image with basic settings
            img_bytes = fig.to_image(
                format="png",
                engine="kaleido",
                width=800,
                height=400
            )
            print("Created Image inside _create_chart_image")
            return img_bytes
        except Exception as e:
            print(f"Error creating chart image: {str(e)}")
            return None
        
    def _add_driver_table_content(self, table: object, driver_info: Dict):
        """Add content to driver information table."""
        try:
            # Set header row
            header_cells = table.rows[0].cells
            headers = ['Dependent Variable', 'Independent Variable', 'Lag', 'Direction']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                #self._apply_cell_style(header_cells[i], is_header=True)

            # Set data row
            data_cells = table.rows[1].cells
            data_cells[0].text = driver_info['dependent_var']
            data_cells[1].text = driver_info['independent_var']
            data_cells[2].text = str(driver_info['lag'])
            data_cells[3].text = driver_info['direction']

            #for cell in data_cells:
                #self._apply_cell_style(cell, is_header=False)

        except Exception as e:
            st.error(f"Error adding driver table content: {str(e)}")

    def _apply_cell_style(self, cell: object, is_header: bool):
        """Apply style to table cell."""
        try:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            font = run.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            if is_header:
                font.bold = True
                font.color.rgb = RGBColor(255, 255, 255)
                cell._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="0000FF"/>')
                )
            else:
                font.bold = False
                font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:
            st.error(f"Error applying cell style: {str(e)}")

    def load_cached_data(self, json_url: str) -> bool:
        """Load cached data from S3 JSON file."""
        try:
            content = self.s3_service.get_document_content(json_url)
            if content:
                self.cached_data = json.loads(content.decode('utf-8'))
                return True
            return False

        except Exception as e:
            st.error(f"Error loading cached data: {str(e)}")
            return False